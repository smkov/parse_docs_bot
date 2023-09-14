# -*- coding: utf-8 -*-
import asyncio
import os
import threading
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup
import logging
from typing import List
from aiogram import Dispatcher, Bot, types
from aiogram_media_group import media_group_handler
from aiohttp import web
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.utils import executor
from settings import *
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram_media_group import media_group_handler
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters import MediaGroupFilter
from docx import Document
from io import BufferedReader
import PyPDF2
import pandas as pd

logging.basicConfig(level=logging.INFO)
bot = Bot(token=dev_bot)
dp = Dispatcher(bot, storage=MemoryStorage())


@dp.message_handler(commands=['start'], state='*')
async def start(message: types.Message):
    await message.answer('Привет, пришли файл\n\nпонимаю .doc .docx .pdf .xlsx')



@dp.message_handler(content_types='document')
async def start(message: types.Message):
    try:
        text = ''
        file = await message.document.download('/')
        if 'doc' in str(file.name).split('.')[-1]:
            document = Document(file.name)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += ' ' + cell.text
            for p in document.paragraphs:
                text += ' ' + p.text
            await bot.send_message(chat_id=message.from_id, text=text)


        if 'pdf' in str(file.name).split('.')[-1]:
            reader = PyPDF2.PdfReader(file.name)
            await bot.send_message(chat_id=message.from_id, text=reader.pages[0].extract_text())
        if 'xlsx' in str(file.name).split('.')[-1]:
            df = pd.read_excel(file.name)
            await bot.send_message(chat_id=message.from_id, text=str(df))

        await asyncio.sleep(1)
        file.close()
        os.remove(file.name)

    except Exception as e:
        print(e)


if __name__ == '__main__':
    try:
        executor.start_polling(dispatcher=dp, skip_updates=True)
    except Exception as e:
        print(e)
